"""Generate the laptop procurement report as a .docx file."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PRIMARY = RGBColor(0x1F, 0x3A, 0x68)   # deep blue
ACCENT = RGBColor(0x2E, 0x75, 0xB6)    # mid blue
LIGHT = RGBColor(0xD9, 0xE2, 0xF3)     # light blue tint
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x59, 0x59, 0x59)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFBFBF", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    if level == 0:
        run.font.size = Pt(22)
        run.font.color.rgb = PRIMARY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = PRIMARY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = ACCENT
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = DARK_TEXT
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=DARK_TEXT, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        run = p.add_run(bold_lead)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.name = "Calibri"
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        set_cell_shading(hdr[i], "1F3A68")
        set_cell_borders(hdr[i], color="1F3A68", size="6")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cells[i], "F2F5FA")
            set_cell_borders(cells[i])

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    return table


def build_report(out_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    add_para(doc, "INTERNAL REPORT TO THE BOARD", bold=True, size=10,
             color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, "Recommended Laptops for Data Science, AI/ML and Analytics Teams", level=0)
    add_para(doc,
             "Procurement proposal for high-performance computing equipment to support secure, "
             "fully local AI, machine learning, and advanced analytics workloads across the institution.",
             italic=True, size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"Prepared: April 2026", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
             "This report recommends a fleet of high-performance laptops to equip seven analytical "
             "teams with the compute capacity required to run modern AI, machine learning and data "
             "analytics workloads on-premises. The recommendations match each team's analytical "
             "focus to a specific machine class so that the institution avoids both under-provisioning "
             "(which slows critical work) and over-provisioning (which wastes capital).")
    add_para(doc,
             "A total of 41 laptops are required across the seven teams. The proposal favours "
             "NVIDIA RTX-class machines for teams that train models locally, Apple Silicon devices "
             "for memory-heavy prototyping and large language model inference, and balanced premium "
             "ultrabooks for teams whose workloads are dominated by data preparation, reporting and "
             "policy analysis.")

    # 2. Background and Objectives
    add_heading(doc, "2. Background and Objectives", level=1)
    add_para(doc,
             "The objective is to acquire high-performance laptops capable of supporting secure, "
             "fully local artificial intelligence, machine learning, and data science workloads "
             "within the institution. These systems will enable:")
    for item in [
        "Large-scale tax data analytics and anomaly detection.",
        "Advanced natural language processing tasks, including large language model (LLM) fine-tuning and inference.",
        "Computer vision applications such as OCR, electronic billing machine (EBM) invoice analysis, and fraud detection.",
        "Multimodal AI systems incorporating retrieval-augmented generation (RAG) over complex document collections containing structured schemas, diagrams, and tables.",
        "Intelligent chatbots and other AI-driven institutional tools built on locally-hosted LLMs and multimodal models.",
    ]:
        add_bullet(doc, item)
    add_para(doc,
             "By processing all data locally, this approach guarantees full data privacy, "
             "eliminates reliance on external cloud GPU services, and provides a scalable "
             "foundation for multiple teams and the future expansion of AI capabilities.")

    # 3. Team Compute Demand Overview
    add_heading(doc, "3. Team Compute Demand Overview", level=1)
    add_para(doc,
             "The table below summarises each team, its analytical focus, and the number of "
             "machines required. Compute intensity is classified as High (sustained local model "
             "training and large-model inference), Medium (regular GPU experimentation, heavy data "
             "engineering or simulation), or Standard (analytics, reporting and lighter modelling).")

    add_table(doc,
              headers=["No.", "Team / Unit", "Core Analytical Focus", "Units", "Compute Intensity"],
              rows=[
                  ["1", "Data Science", "AI/ML models, big data analytics, automation", "11", "High"],
                  ["2", "Revenue Modeling and Statistics", "Econometric modeling, microsimulation, tax analysis", "6", "Medium"],
                  ["3", "Tax Analysis & Revenue Forecasting", "Time-series forecasting, scenario analysis", "4", "Medium"],
                  ["4", "Data Governance", "Data quality, metadata management, system integration", "4", "Standard"],
                  ["5", "Research and Policy Analysis", "Impact evaluation, policy simulations, analytical research", "8", "Medium"],
                  ["6", "EBM Data Analysis", "Evidence-based monitoring, dashboards, performance analysis", "5", "Standard"],
                  ["7", "Statistics Customs", "Customs data analysis", "3", "Standard"],
                  ["", "Total", "", "41", ""],
              ],
              col_widths=[Cm(1.0), Cm(4.5), Cm(6.5), Cm(1.6), Cm(3.2)])

    # 4. Selection Methodology
    add_heading(doc, "4. Selection Methodology", level=1)
    add_para(doc,
             "Recommendations were made by mapping each team's daily workload to four buying "
             "criteria that drive real-world performance in AI and data science: GPU capability "
             "and VRAM, RAM capacity, SSD speed and capacity, and sustained thermal performance. "
             "Peak benchmark scores were de-prioritised in favour of sustained throughput, because "
             "training and large-data workflows expose throttling far more than short bursts.")

    add_para(doc, "Three machine tiers were defined:", bold=True)
    add_table(doc,
              headers=["Tier", "Profile", "Target RAM", "Target Storage", "GPU Guidance"],
              rows=[
                  ["Pro", "Local deep-learning training, large-model inference, multi-project workflows",
                   "64–128 GB", "2 TB NVMe", "RTX 5090 (24 GB) or Apple M5 Max"],
                  ["Recommended", "Mixed ML, simulation, heavy data engineering, GPU-accelerated experiments",
                   "32–64 GB", "1–2 TB NVMe", "RTX 5070 Ti / 5080 (12–16 GB)"],
                  ["Standard", "Analytics, dashboards, policy analysis, governance and reporting",
                   "32 GB", "1 TB NVMe", "Integrated or entry RTX"],
              ],
              col_widths=[Cm(2.5), Cm(6.5), Cm(2.8), Cm(2.8), Cm(4.5)])

    # 5. Team-by-team recommendations
    add_heading(doc, "5. Team-by-Team Recommendations", level=1)

    # 5.1 Data Science
    add_heading(doc, "5.1 Data Science (11 units) — Pro tier", level=2)
    add_para(doc,
             "This team carries the institution's heaviest AI workload: training and fine-tuning "
             "machine-learning models, running large language models locally, performing big-data "
             "analytics, and automating intelligence pipelines. VRAM, sustained GPU power and "
             "memory bandwidth are the binding constraints, and any under-provisioning here "
             "directly limits what the institution can build in-house.")
    add_para(doc, "Primary recommendation: ASUS ROG Zephyrus Duo (2026) GX651 — 8 units", bold=True)
    add_bullet(doc, "Intel Core Ultra 9 386H (16 cores) with NPU up to 50 TOPS for on-device acceleration.")
    add_bullet(doc, "NVIDIA RTX 5090 Laptop GPU with 24 GB GDDR7 — sufficient for fine-tuning mid-sized transformers and running 13B–34B parameter LLMs locally.")
    add_bullet(doc, "64 GB LPDDR5X memory and 2 TB PCIe 5.0 NVMe SSD for large dataset handling and rapid checkpointing.")
    add_bullet(doc, "Dual-screen ROG Nebula HDR OLED display improves productivity in notebooks, dashboards, and model monitoring.")

    add_para(doc, "Complementary recommendation: MacBook Pro 16-inch (M5 Max) — 3 units", bold=True)
    add_bullet(doc, "Up to 128 GB unified memory enables loading of 70B-parameter models for inference and RAG experimentation, which exceeds the practical reach of any 24 GB VRAM laptop.")
    add_bullet(doc, "16-core Neural Engine and 40-core GPU accelerate transformer workloads via PyTorch's MPS backend.")
    add_bullet(doc, "Up to 24 hours battery life and silent operation suits team leads who travel between sites and conduct on-the-go research.")
    add_para(doc,
             "Justification for the split: a small number of M5 Max machines extends the team's "
             "ceiling for very large model inference (where unified memory matters more than CUDA), "
             "while the majority RTX fleet retains full compatibility with CUDA-based training "
             "stacks used by the open-source ML ecosystem.", italic=True, color=GREY)

    # 5.2 Revenue Modeling and Statistics
    add_heading(doc, "5.2 Revenue Modeling and Statistics (6 units) — Recommended tier", level=2)
    add_para(doc,
             "This team performs econometric modelling, microsimulation and tax analysis. These "
             "workloads are CPU- and memory-bound — long simulation runs, large panel datasets "
             "and Monte Carlo iterations benefit far more from sustained CPU performance and "
             "ample RAM than from peak GPU power.")
    add_para(doc, "Recommendation: Lenovo Legion Pro 7i Gen 10 (16-inch) — 6 units", bold=True)
    add_bullet(doc, "Intel Core Ultra 9 / RTX 5080 Laptop GPU (16 GB VRAM) — provides accelerated GPU capability for matrix-heavy simulations without the cost of an RTX 5090 fleet.")
    add_bullet(doc, "64 GB DDR5 RAM and 2 TB NVMe SSD support large microsimulation datasets and multi-year tax panels held entirely in memory.")
    add_bullet(doc, "Strong cooling system sustains long simulation runs without thermal throttling.")

    # 5.3 Tax Analysis & Revenue Forecasting
    add_heading(doc, "5.3 Tax Analysis & Revenue Forecasting (4 units) — Recommended tier", level=2)
    add_para(doc,
             "This team produces time-series forecasts and scenario analyses that feed directly "
             "into revenue planning. Workloads include statistical models, gradient-boosted "
             "ensembles and neural forecasting architectures that benefit from a moderate GPU "
             "and a generous RAM allocation.")
    add_para(doc, "Recommendation: ASUS ROG Zephyrus G14 (2025) — 4 units", bold=True)
    add_bullet(doc, "NVIDIA RTX 5070 Ti Laptop GPU — accelerates training of forecasting models (LSTM, Transformer-based forecasters) without the bulk of a desktop replacement.")
    add_bullet(doc, "32 GB DDR5 RAM and 1 TB NVMe SSD — comfortable for parallel scenario runs and concurrent notebook sessions.")
    add_bullet(doc, "Compact 14-inch chassis suits analysts who frequently present scenarios to leadership and policy stakeholders.")

    # 5.4 Data Governance
    add_heading(doc, "5.4 Data Governance (4 units) — Standard tier", level=2)
    add_para(doc,
             "The team's focus is data quality, metadata management and system integration. "
             "Workloads are dominated by SQL, ETL orchestration, schema validation and integration "
             "tooling, where stability, port selection and a reliable enterprise-grade build "
             "matter more than GPU power.")
    add_para(doc, "Recommendation: Lenovo ThinkPad P1 Gen 7 — 4 units", bold=True)
    add_bullet(doc, "Workstation-grade reliability with certified drivers — important for teams that interface with multiple institutional systems and docks.")
    add_bullet(doc, "32 GB DDR5 RAM and 1 TB NVMe SSD — supports multiple database clients, virtual machines and data quality tooling running simultaneously.")
    add_bullet(doc, "Entry RTX option provides headroom for occasional analytical workloads without paying for a gaming-class GPU.")

    # 5.5 Research and Policy Analysis
    add_heading(doc, "5.5 Research and Policy Analysis (8 units) — Recommended tier", level=2)
    add_para(doc,
             "This team conducts impact evaluations, policy simulations and analytical research, "
             "frequently combining econometric techniques with policy-document review and "
             "presentation work. The workload is mixed: long simulation runs, heavy reading, "
             "and significant time spent producing publication-quality outputs.")
    add_para(doc, "Recommendation: Dell XPS 16 (9640) — 8 units", bold=True)
    add_bullet(doc, "Premium 16-inch OLED display — reduces fatigue during the long reading and writing sessions characteristic of policy work.")
    add_bullet(doc, "32 GB DDR5 RAM, 1 TB NVMe SSD, and an RTX 4060 GPU option — sufficient for policy simulations and light-to-moderate ML experimentation.")
    add_bullet(doc, "Balanced ergonomics and portability for researchers presenting to government and external partners.")

    # 5.6 EBM Data Analysis
    add_heading(doc, "5.6 EBM Data Analysis (5 units) — Standard tier", level=2)
    add_para(doc,
             "This team performs evidence-based monitoring, builds dashboards, and analyses "
             "performance data drawn from the EBM ecosystem. Workloads are I/O- and "
             "visualisation-heavy rather than compute-bound: BI tools, SQL clients and dashboard "
             "platforms benefit most from RAM, fast storage and a high-quality display.")
    add_para(doc, "Recommendation: Dell XPS 16 (9640) — 5 units", bold=True)
    add_bullet(doc, "32 GB DDR5 RAM and 1 TB NVMe SSD — keeps multiple dashboards, BI tools and large query results responsive.")
    add_bullet(doc, "OLED display with full DCI-P3 colour coverage produces accurate, presentation-ready dashboards and visualisations.")
    add_bullet(doc, "Standardising EBM and Research and Policy Analysis on the same model simplifies support, imaging and spare-part inventory.")

    # 5.7 Statistics Customs
    add_heading(doc, "5.7 Statistics Customs (3 units) — Standard tier", level=2)
    add_para(doc,
             "Customs statistics work centres on transactional and tariff data analysis, with a "
             "focus on accurate reporting rather than model training. A reliable, value-oriented "
             "machine with strong CPU and ample RAM is the right fit.")
    add_para(doc, "Recommendation: ASUS TUF Gaming A16 (2025) — 3 units", bold=True)
    add_bullet(doc, "RTX 5060 Laptop GPU with 32 GB DDR5 RAM and 1 TB NVMe SSD — provides GPU headroom for occasional ML experimentation while keeping unit cost contained.")
    add_bullet(doc, "Stronger cooling profile than thin-and-light alternatives supports sustained data-processing tasks on large customs datasets.")
    add_bullet(doc, "Best price-to-performance ratio in the fleet, allowing budget to be redirected to higher tiers where it is more impactful.")

    # 6. Consolidated procurement table
    add_heading(doc, "6. Consolidated Procurement Plan", level=1)
    add_para(doc,
             "The table below consolidates the recommended models, quantities and key configuration "
             "targets across the seven teams. Final unit pricing should be confirmed with authorised "
             "regional vendors at the point of order.")

    add_table(doc,
              headers=["Team / Unit", "Recommended Model", "Tier", "Qty", "Key Configuration"],
              rows=[
                  ["Data Science", "ASUS ROG Zephyrus Duo (2026) GX651", "Pro", "8",
                   "RTX 5090 24 GB · 64 GB RAM · 2 TB NVMe"],
                  ["Data Science", "MacBook Pro 16\" (M5 Max)", "Pro", "3",
                   "M5 Max · 128 GB unified · 2 TB SSD"],
                  ["Revenue Modeling and Statistics", "Lenovo Legion Pro 7i Gen 10", "Recommended", "6",
                   "RTX 5080 16 GB · 64 GB RAM · 2 TB NVMe"],
                  ["Tax Analysis & Revenue Forecasting", "ASUS ROG Zephyrus G14 (2025)", "Recommended", "4",
                   "RTX 5070 Ti · 32 GB RAM · 1 TB NVMe"],
                  ["Data Governance", "Lenovo ThinkPad P1 Gen 7", "Standard", "4",
                   "Entry RTX · 32 GB RAM · 1 TB NVMe"],
                  ["Research and Policy Analysis", "Dell XPS 16 (9640)", "Recommended", "8",
                   "RTX 4060 · 32 GB RAM · 1 TB NVMe"],
                  ["EBM Data Analysis", "Dell XPS 16 (9640)", "Standard", "5",
                   "Integrated/RTX 4050 · 32 GB RAM · 1 TB NVMe"],
                  ["Statistics Customs", "ASUS TUF Gaming A16 (2025)", "Standard", "3",
                   "RTX 5060 · 32 GB RAM · 1 TB NVMe"],
                  ["Total", "", "", "41", ""],
              ],
              col_widths=[Cm(4.5), Cm(5.0), Cm(2.4), Cm(1.2), Cm(5.5)])

    # 7. Risk and mitigation
    add_heading(doc, "7. Key Risks and Mitigations", level=1)
    add_bullet(doc, "Soldered memory in some Pro-tier models limits future upgrades.",
               bold_lead="Memory ceiling: ")
    add_para(doc,
             "Mitigation: machines are specified at their highest practical RAM at purchase, so no "
             "upgrade is required within the planned lifecycle.", size=10.5, color=GREY)
    add_bullet(doc, "Sustained training runs generate significant heat in compact laptops.",
               bold_lead="Thermal throttling: ")
    add_para(doc,
             "Mitigation: Pro-tier units are deliberately specified on chassis with strong cooling "
             "(Zephyrus Duo, Legion Pro 7i) rather than the thinnest available designs.",
             size=10.5, color=GREY)
    add_bullet(doc, "Apple Silicon does not support CUDA; some open-source training pipelines assume CUDA availability.",
               bold_lead="CUDA compatibility: ")
    add_para(doc,
             "Mitigation: M5 Max units are allocated only to use cases dominated by inference and "
             "memory-bound workloads, where the MPS backend is sufficient.",
             size=10.5, color=GREY)
    add_bullet(doc, "Premium GPU laptops carry higher unit cost than office-grade machines.",
               bold_lead="Cost concentration: ")
    add_para(doc,
             "Mitigation: tiering ensures premium specifications are reserved for the Data Science "
             "team, where the institutional return on AI capability is highest. Lower-intensity "
             "teams receive appropriately scaled machines.",
             size=10.5, color=GREY)

    # 8. Recommendation
    add_heading(doc, "8. Recommendation to the Board", level=1)
    add_para(doc,
             "The Board is invited to approve the procurement of the 41 laptops detailed in "
             "Section 6, distributed across the seven teams in line with the tiering and "
             "justifications presented in Section 5. The proposed fleet provides the institution "
             "with a defensible, secure and fully local AI compute capability while keeping "
             "investment proportional to each team's demonstrated workload.")
    add_para(doc,
             "Approving this proposal will allow the institution to scale its in-house AI, machine "
             "learning and analytics capabilities without dependence on external cloud GPU services, "
             "preserving full data sovereignty over sensitive tax, customs and policy data.",
             italic=True, color=GREY)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_report("/Users/blaiseai4sense/Desktop/ClaudeRemyProj/Laptop_Procurement_Report.docx")
